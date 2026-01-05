"""
File Processing Services
Handles extraction of text and metadata from various file formats
"""

import os
import io
from typing import Tuple, Dict
import mimetypes

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word processing
try:
    from docx import Document
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False

# Excel processing
try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Image OCR
try:
    from google.cloud import vision
    import PIL.Image
    VISION_AVAILABLE = True
except ImportError:
    VISION_AVAILABLE = False


async def process_pdf(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from PDF files
    Returns: (extracted_text, metadata)
    """
    if not PDF_AVAILABLE:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    
    try:
        extracted_text = ""
        metadata = {
            "pages": 0,
            "images": 0,
            "tables": 0
        }
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(pdf_reader.pages)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n"
                        extracted_text += page_text
                except Exception as e:
                    print(f"Error extracting page {page_num + 1}: {e}")
            
            # Get PDF metadata if available
            if pdf_reader.metadata:
                if pdf_reader.metadata.get('/Title'):
                    metadata["title"] = pdf_reader.metadata.get('/Title')
                if pdf_reader.metadata.get('/Author'):
                    metadata["author"] = pdf_reader.metadata.get('/Author')
                if pdf_reader.metadata.get('/Subject'):
                    metadata["subject"] = pdf_reader.metadata.get('/Subject')
        
        return extracted_text.strip(), metadata
        
    except Exception as e:
        raise Exception(f"PDF processing failed: {str(e)}")


async def process_word(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from Word documents (.docx)
    Returns: (extracted_text, metadata)
    """
    if not WORD_AVAILABLE:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    try:
        doc = Document(file_path)
        
        extracted_text = ""
        metadata = {
            "paragraphs": 0,
            "tables": 0,
            "images": 0
        }
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"
                metadata["paragraphs"] += 1
        
        # Extract tables
        for table in doc.tables:
            metadata["tables"] += 1
            extracted_text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                extracted_text += row_text + "\n"
        
        # Count images
        metadata["images"] = len([r for r in doc.part.rels.values() 
                                  if "image" in r.target_ref])
        
        # Get core properties
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            if doc.core_properties.subject:
                metadata["subject"] = doc.core_properties.subject
        
        return extracted_text.strip(), metadata
        
    except Exception as e:
        raise Exception(f"Word processing failed: {str(e)}")


async def process_excel(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from Excel files (.xlsx, .xls)
    Returns: (extracted_text, metadata)
    """
    if not EXCEL_AVAILABLE:
        raise ImportError("openpyxl and pandas not installed. Run: pip install openpyxl pandas")
    
    try:
        # Load workbook
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        extracted_text = ""
        metadata = {
            "sheets": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
            "total_rows": 0,
            "total_columns": 0
        }
        
        # Process each sheet
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            extracted_text += f"\n=== Sheet: {sheet_name} ===\n"
            
            # Convert to pandas for easier processing
            data = ws.values
            cols = next(data)
            data_list = list(data)
            
            if data_list:
                df = pd.DataFrame(data_list, columns=cols)
                
                # Remove completely empty rows and columns
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                metadata["total_rows"] += len(df)
                metadata["total_columns"] = max(metadata["total_columns"], len(df.columns))
                
                # Convert to text
                extracted_text += df.to_string(index=False, na_rep='')
                extracted_text += "\n"
        
        return extracted_text.strip(), metadata
        
    except Exception as e:
        raise Exception(f"Excel processing failed: {str(e)}")


async def process_image(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from images using Google Vision API (OCR)
    Falls back to basic image info if Vision API not available
    Returns: (extracted_text, metadata)
    """
    metadata = {
        "ocr_performed": False,
        "width": 0,
        "height": 0,
        "format": ""
    }
    
    try:
        # Get basic image info
        from PIL import Image
        with Image.open(file_path) as img:
            metadata["width"] = img.width
            metadata["height"] = img.height
            metadata["format"] = img.format
        
        # Attempt OCR with Google Vision
        if VISION_AVAILABLE and os.getenv('GOOGLE_APPLICATION_CREDENTIALS'):
            client = vision.ImageAnnotatorClient()
            
            with open(file_path, 'rb') as image_file:
                content = image_file.read()
            
            image = vision.Image(content=content)
            response = client.text_detection(image=image)
            texts = response.text_annotations
            
            if texts:
                extracted_text = texts[0].description
                metadata["ocr_performed"] = True
                metadata["confidence"] = "high"  # Google Vision is generally high confidence
                return extracted_text.strip(), metadata
        
        # Fallback: Return basic info
        extracted_text = f"Image file: {os.path.basename(file_path)}\n"
        extracted_text += f"Dimensions: {metadata['width']}x{metadata['height']}\n"
        extracted_text += f"Format: {metadata['format']}\n"
        extracted_text += "Note: OCR not performed (Google Vision API not configured)"
        
        return extracted_text, metadata
        
    except Exception as e:
        raise Exception(f"Image processing failed: {str(e)}")


async def process_text(file_path: str) -> Tuple[str, Dict]:
    """
    Extract text from plain text files
    Returns: (extracted_text, metadata)
    """
    try:
        # Try to detect encoding
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        extracted_text = None
        encoding_used = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    extracted_text = f.read()
                    encoding_used = encoding
                    break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if extracted_text is None:
            raise Exception("Could not decode text file with any known encoding")
        
        metadata = {
            "encoding": encoding_used,
            "lines": len(extracted_text.split('\n')),
            "characters": len(extracted_text),
            "words": len(extracted_text.split())
        }
        
        return extracted_text.strip(), metadata
        
    except Exception as e:
        raise Exception(f"Text processing failed: {str(e)}")


async def process_file_auto(file_path: str) -> Tuple[str, Dict]:
    """
    Automatically detect file type and process accordingly
    Returns: (extracted_text, metadata)
    """
    mime_type, _ = mimetypes.guess_type(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    if not mime_type:
        # Try to guess from extension
        mime_type_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain',
            '.png': 'image/png',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg'
        }
        mime_type = mime_type_map.get(ext)
    
    if not mime_type:
        raise Exception(f"Could not determine file type for: {file_path}")
    
    # Determine file_type for database constraint
    file_type_map = {
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.xlsx': 'excel',
        '.xls': 'excel',
        '.txt': 'text',
        '.png': 'image',
        '.jpg': 'image',
        '.jpeg': 'image'
    }
    file_type = file_type_map.get(ext, 'unknown')
    
    # Route to appropriate processor
    if 'pdf' in mime_type:
        text, metadata = await process_pdf(file_path)
    elif 'wordprocessingml' in mime_type or 'msword' in mime_type:
        # IMPORTANT:
        # python-docx can only read modern .docx files (ZIP package).
        # Old binary .doc files will fail with "Package not found".
        if ext == '.doc':
            raise Exception(
                "Unsupported Word format: .doc (old Word). Please open it in Word/Google Docs and 'Save As' .docx, "
                "or export as PDF, then upload again."
            )
        text, metadata = await process_word(file_path)
    elif 'spreadsheetml' in mime_type or 'ms-excel' in mime_type:
        text, metadata = await process_excel(file_path)
    elif 'image' in mime_type:
        text, metadata = await process_image(file_path)
    elif 'text' in mime_type:
        text, metadata = await process_text(file_path)
    else:
        raise Exception(f"Unsupported file type: {mime_type}")
    
    # Add file_type to metadata
    metadata['file_type'] = file_type
    return text, metadata


# Utility function for checking dependencies
def check_dependencies() -> Dict[str, bool]:
    """Check which file processing libraries are available"""
    return {
        "pdf": PDF_AVAILABLE,
        "word": WORD_AVAILABLE,
        "excel": EXCEL_AVAILABLE,
        "ocr": VISION_AVAILABLE,
        "all": all([PDF_AVAILABLE, WORD_AVAILABLE, EXCEL_AVAILABLE])
    }


# Installation helper
def get_installation_instructions() -> str:
    """Get installation instructions for missing dependencies"""
    deps = check_dependencies()
    instructions = []
    
    if not deps["pdf"]:
        instructions.append("PDF support: pip install PyPDF2")
    if not deps["word"]:
        instructions.append("Word support: pip install python-docx")
    if not deps["excel"]:
        instructions.append("Excel support: pip install openpyxl pandas")
    if not deps["ocr"]:
        instructions.append("OCR support: pip install google-cloud-vision pillow")
        instructions.append("  Also set GOOGLE_APPLICATION_CREDENTIALS environment variable")
    
    if instructions:
        return "Missing dependencies:\n" + "\n".join(instructions)
    else:
        return "All dependencies installed!"
